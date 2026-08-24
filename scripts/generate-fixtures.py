#!/usr/bin/env python3
"""Render the versioned synthetic source corpus into upload-ready fixtures."""

from __future__ import annotations

import hashlib
import json
import re
from dataclasses import dataclass
from html import escape
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from pypdf import PdfReader
from reportlab.lib.colors import HexColor
from reportlab.lib.enums import TA_LEFT
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import (
    BaseDocTemplate,
    Flowable,
    Frame,
    PageBreak,
    PageTemplate,
    Paragraph,
)

REPOSITORY_ROOT = Path(__file__).resolve().parents[1]
SOURCE_MANIFEST = REPOSITORY_ROOT / "evals" / "dataset" / "source-manifest.json"
OUTPUT_ROOT = REPOSITORY_ROOT / "fixtures" / "documents"
GENERATED_MANIFEST = OUTPUT_ROOT / "manifest.json"
MARKER_PATTERN = re.compile(
    r"^\[(?P<marker>LG-(?:POL|ATK)-\d{3}:(?P<kind>H\d{2}|P\d{3}|L\d{3}))\]\s*(?P<text>.*)$"
)
HEADER_PATTERN = re.compile(r"^(?P<key>[A-Z_]+):\s*(?P<value>.*)$")

NAVY = "0F2747"
BLUE = "2E74B5"
TEAL = "087E8B"
SLATE = "475569"
LIGHT_SLATE = "E2E8F0"


@dataclass(frozen=True)
class SourceDocument:
    source_id: str
    title: str
    target_format: str
    kind: str
    source_path: Path
    headers: dict[str, str]
    lines: list[tuple[str | None, str | None, str]]


class MarkerParagraph(Paragraph):
    """ReportLab paragraph carrying a stable source marker."""

    def __init__(self, *args: Any, marker_id: str | None = None, **kwargs: Any) -> None:
        super().__init__(*args, **kwargs)
        self.marker_id = marker_id


class FixturePdfTemplate(BaseDocTemplate):
    """PDF template that records the actual page of each source marker."""

    def __init__(self, filename: str, *, source_id: str, title: str) -> None:
        super().__init__(
            filename,
            pagesize=LETTER,
            leftMargin=inch,
            rightMargin=inch,
            topMargin=0.82 * inch,
            bottomMargin=0.72 * inch,
            title=title,
            author="LocalGuard Demonstration Organization",
            subject="Original synthetic LocalGuard AI evaluation fixture",
        )
        self.source_id = source_id
        self.fixture_title = title
        self.marker_pages: dict[str, int] = {}
        frame = Frame(
            self.leftMargin,
            self.bottomMargin,
            self.width,
            self.height,
            id="fixture",
            leftPadding=0,
            rightPadding=0,
            topPadding=0,
            bottomPadding=0,
        )
        self.addPageTemplates(PageTemplate(id="policy", frames=[frame], onPage=self._draw_page))

    def _draw_page(self, canvas: Any, document: Any) -> None:
        canvas.saveState()
        canvas.setStrokeColor(HexColor(f"#{TEAL}"))
        canvas.setLineWidth(2)
        canvas.line(inch, 10.38 * inch, 7.5 * inch, 10.38 * inch)
        canvas.setFillColor(HexColor(f"#{SLATE}"))
        canvas.setFont("Helvetica", 8)
        canvas.drawString(inch, 10.52 * inch, f"{self.source_id}  |  SYNTHETIC FIXTURE")
        canvas.drawRightString(7.5 * inch, 0.42 * inch, f"Page {document.page}")
        canvas.drawString(inch, 0.42 * inch, "LocalGuard Demonstration Organization")
        canvas.restoreState()

    def afterFlowable(self, flowable: Flowable) -> None:
        marker_id = getattr(flowable, "marker_id", None)
        if marker_id:
            self.marker_pages[marker_id] = self.page


def _sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as stream:
        for block in iter(lambda: stream.read(1024 * 1024), b""):
            digest.update(block)
    return digest.hexdigest()


def _read_source(entry: dict[str, Any]) -> SourceDocument:
    source_path = REPOSITORY_ROOT / entry["path"]
    raw_lines = source_path.read_text(encoding="utf-8").splitlines()
    headers: dict[str, str] = {}
    parsed_lines: list[tuple[str | None, str | None, str]] = []
    for raw_line in raw_lines:
        marker_match = MARKER_PATTERN.match(raw_line)
        if marker_match:
            parsed_lines.append(
                (
                    marker_match.group("marker"),
                    marker_match.group("kind")[0],
                    marker_match.group("text"),
                )
            )
            continue
        header_match = HEADER_PATTERN.match(raw_line)
        if header_match:
            headers[header_match.group("key")] = header_match.group("value")
        elif raw_line.strip():
            parsed_lines.append((None, None, raw_line.strip()))
    return SourceDocument(
        source_id=entry["source_id"],
        title=headers["TITLE"],
        target_format=entry["target_format"],
        kind=entry["kind"],
        source_path=source_path,
        headers=headers,
        lines=parsed_lines,
    )


def _pdf_styles() -> dict[str, ParagraphStyle]:
    base = getSampleStyleSheet()
    return {
        "title": ParagraphStyle(
            "FixtureTitle",
            parent=base["Title"],
            fontName="Helvetica-Bold",
            fontSize=23,
            leading=27,
            textColor=HexColor(f"#{NAVY}"),
            alignment=TA_LEFT,
            spaceBefore=0,
            spaceAfter=10,
        ),
        "subtitle": ParagraphStyle(
            "FixtureSubtitle",
            parent=base["Normal"],
            fontName="Helvetica",
            fontSize=9,
            leading=12,
            textColor=HexColor(f"#{SLATE}"),
            spaceAfter=4,
        ),
        "heading": ParagraphStyle(
            "FixtureHeading",
            parent=base["Heading1"],
            fontName="Helvetica-Bold",
            fontSize=15,
            leading=19,
            textColor=HexColor(f"#{BLUE}"),
            spaceBefore=14,
            spaceAfter=7,
            keepWithNext=True,
        ),
        "body": ParagraphStyle(
            "FixtureBody",
            parent=base["BodyText"],
            fontName="Helvetica",
            fontSize=10.5,
            leading=14,
            textColor=HexColor(f"#{NAVY}"),
            spaceAfter=7,
        ),
        "marker": ParagraphStyle(
            "FixtureMarker",
            parent=base["Normal"],
            fontName="Courier",
            fontSize=7,
            leading=9,
            textColor=HexColor(f"#{SLATE}"),
            spaceAfter=2,
        ),
        "notice": ParagraphStyle(
            "FixtureNotice",
            parent=base["Normal"],
            fontName="Helvetica-Oblique",
            fontSize=8.5,
            leading=11,
            textColor=HexColor(f"#{SLATE}"),
            borderColor=HexColor(f"#{LIGHT_SLATE}"),
            borderWidth=0.5,
            borderPadding=7,
            backColor=HexColor("#F8FAFC"),
            spaceBefore=7,
            spaceAfter=12,
        ),
    }


def _render_pdf(source: SourceDocument, target: Path) -> dict[str, dict[str, Any]]:
    styles = _pdf_styles()
    source_id = escape(source.source_id)
    version = escape(source.headers["VERSION"])
    document_class = escape(source.headers["DOCUMENT_CLASS"])
    story: list[Flowable] = [
        Paragraph(escape(source.title), styles["title"]),
        Paragraph(
            f"<b>{source_id}</b> &nbsp; Version {version} &nbsp; Class: {document_class}",
            styles["subtitle"],
        ),
        Paragraph(escape(source.headers["SYNTHETIC_NOTICE"]), styles["notice"]),
    ]
    heading_total = sum(1 for _, kind, _ in source.lines if kind == "H")
    second_page_heading = max(2, (heading_total // 2) + 1)
    heading_number = 0
    for marker_id, kind, text in source.lines:
        if kind == "H":
            heading_number += 1
            if heading_number == second_page_heading:
                story.append(PageBreak())
            marker = escape(marker_id or "")
            content = f'<font name="Courier" size="7" color="#{SLATE}">[{marker}]</font><br/>'
            content += escape(text)
            story.append(MarkerParagraph(content, styles["heading"], marker_id=marker_id))
        elif kind == "P":
            story.append(
                MarkerParagraph(
                    f"[{escape(marker_id or '')}]",
                    styles["marker"],
                    marker_id=marker_id,
                )
            )
        elif kind == "L":
            marker = escape(marker_id or "")
            content = f'<font name="Courier" size="7" color="#{SLATE}">[{marker}]</font> '
            content += escape(text)
            story.append(MarkerParagraph(content, styles["body"], marker_id=marker_id))
        elif text:
            story.append(Paragraph(escape(text), styles["body"]))
    template = FixturePdfTemplate(str(target), source_id=source.source_id, title=source.title)
    template.build(story)
    page_count = len(PdfReader(str(target)).pages)
    return {
        marker: {"anchor_type": "page", "page": page}
        for marker, page in sorted(template.marker_pages.items())
    } | {"_document": {"page_count": page_count}}


def _set_run_font(
    run: Any,
    *,
    name: str = "Calibri",
    size: float | None = None,
    color: str | None = None,
    bold: bool | None = None,
    italic: bool | None = None,
) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def _configure_docx(document: Document) -> None:
    section = document.sections[0]
    section.start_type = WD_SECTION.NEW_PAGE
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    heading = document.styles["Heading 1"]
    heading.font.name = "Calibri"
    heading._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    heading._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    heading.font.size = Pt(16)
    heading.font.bold = True
    heading.font.color.rgb = RGBColor.from_string(BLUE)
    heading.paragraph_format.space_before = Pt(18)
    heading.paragraph_format.space_after = Pt(10)
    heading.paragraph_format.keep_with_next = True


def _add_page_field(paragraph: Any) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    _set_run_font(run, size=8, color=SLATE)
    field_begin = OxmlElement("w:fldChar")
    field_begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    field_end = OxmlElement("w:fldChar")
    field_end.set(qn("w:fldCharType"), "end")
    run._r.extend([field_begin, instruction, field_end])


def _render_docx(source: SourceDocument, target: Path) -> dict[str, dict[str, Any]]:
    document = Document()
    _configure_docx(document)
    section = document.sections[0]

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header_run = header.add_run(f"{source.source_id}  |  SYNTHETIC FIXTURE")
    _set_run_font(header_run, size=8, color=SLATE, bold=True)
    _add_page_field(section.footer.paragraphs[0])

    title = document.add_paragraph()
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(8)
    title_run = title.add_run(source.title)
    _set_run_font(title_run, size=23, color=NAVY, bold=True)

    metadata = document.add_paragraph()
    metadata.paragraph_format.space_after = Pt(8)
    metadata_run = metadata.add_run(
        f"{source.source_id} | Version {source.headers['VERSION']} | "
        f"Class: {source.headers['DOCUMENT_CLASS']}"
    )
    _set_run_font(metadata_run, size=9, color=SLATE)

    notice = document.add_paragraph()
    notice.paragraph_format.space_after = Pt(12)
    notice_run = notice.add_run(source.headers["SYNTHETIC_NOTICE"])
    _set_run_font(notice_run, size=8.5, color=SLATE, italic=True)

    marker_locations: dict[str, dict[str, Any]] = {}
    for marker_id, kind, text in source.lines:
        if kind == "H":
            paragraph = document.add_paragraph(style="Heading 1")
            marker_run = paragraph.add_run(f"[{marker_id}] ")
            _set_run_font(marker_run, name="Consolas", size=7, color=SLATE)
            text_run = paragraph.add_run(text)
            _set_run_font(text_run, size=16, color=BLUE, bold=True)
        elif kind == "P":
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(2)
            marker_run = paragraph.add_run(f"[{marker_id}]")
            _set_run_font(marker_run, name="Consolas", size=7, color=SLATE)
        else:
            paragraph = document.add_paragraph()
            marker_run = paragraph.add_run(f"[{marker_id}] " if marker_id else "")
            _set_run_font(marker_run, name="Consolas", size=7, color=SLATE)
            text_run = paragraph.add_run(text)
            _set_run_font(text_run, size=11, color=NAVY)
        if marker_id:
            marker_locations[marker_id] = {
                "anchor_type": "paragraph",
                "paragraph": len(document.paragraphs),
            }

    core = document.core_properties
    core.title = source.title
    core.subject = "Original synthetic LocalGuard AI evaluation fixture"
    core.author = "LocalGuard Demonstration Organization"
    core.comments = "Synthetic content only; no real organization or personal data."
    core.keywords = "LocalGuard AI, synthetic, evaluation fixture"
    core.last_modified_by = "LocalGuard Demonstration Organization"
    document.save(target)
    return marker_locations | {"_document": {"paragraph_count": len(document.paragraphs)}}


def _render_txt(source: SourceDocument, target: Path) -> dict[str, dict[str, Any]]:
    text = source.source_path.read_text(encoding="utf-8").replace("\r\n", "\n")
    target.write_text(text.rstrip() + "\n", encoding="utf-8", newline="\n")
    marker_locations: dict[str, dict[str, Any]] = {}
    for line_number, line in enumerate(text.splitlines(), start=1):
        marker_match = MARKER_PATTERN.match(line)
        if marker_match:
            marker_locations[marker_match.group("marker")] = {
                "anchor_type": "line",
                "line_start": line_number,
                "line_end": line_number,
            }
    return marker_locations | {"_document": {"line_count": len(text.splitlines())}}


def main() -> int:
    manifest = json.loads(SOURCE_MANIFEST.read_text(encoding="utf-8"))
    if not manifest.get("synthetic_only"):
        raise RuntimeError("Fixture generation refuses a corpus not declared synthetic-only.")

    (OUTPUT_ROOT / "clean").mkdir(parents=True, exist_ok=True)
    (OUTPUT_ROOT / "attacks").mkdir(parents=True, exist_ok=True)

    generated: list[dict[str, Any]] = []
    for entry in manifest["sources"]:
        source = _read_source(entry)
        output_directory = OUTPUT_ROOT / ("clean" if source.kind == "clean" else "attacks")
        target = output_directory / f"{source.source_path.stem}.{source.target_format}"
        if source.target_format == "pdf":
            locations = _render_pdf(source, target)
        elif source.target_format == "docx":
            locations = _render_docx(source, target)
        elif source.target_format == "txt":
            locations = _render_txt(source, target)
        else:
            raise ValueError(f"Unsupported target format: {source.target_format}")

        generated.append(
            {
                "source_id": source.source_id,
                "kind": source.kind,
                "format": source.target_format,
                "path": target.relative_to(REPOSITORY_ROOT).as_posix(),
                "source_path": source.source_path.relative_to(REPOSITORY_ROOT).as_posix(),
                "source_sha256": _sha256(source.source_path),
                "sha256": _sha256(target),
                "bytes": target.stat().st_size,
                "locations": locations,
            }
        )

    output_manifest = {
        "dataset_version": manifest["dataset_version"],
        "synthetic_only": True,
        "generator": "scripts/generate-fixtures.py",
        "documents": generated,
    }
    GENERATED_MANIFEST.write_text(
        json.dumps(output_manifest, indent=2, sort_keys=True) + "\n",
        encoding="utf-8",
        newline="\n",
    )
    print(f"Generated {len(generated)} upload-ready synthetic fixtures in {OUTPUT_ROOT}")
    print(f"Manifest: {GENERATED_MANIFEST}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
