from __future__ import annotations

import io
import uuid
import zipfile

import pytest
from docx import Document
from localguard_api.config import Settings
from localguard_api.errors import UnsafeUploadError
from localguard_api.ingestion import (
    ValidatedUpload,
    _decode_text,
    _validate_docx_archive,
    build_chunks,
    parse_document,
)

pytestmark = [pytest.mark.unit, pytest.mark.security]


def _settings(**values: object) -> Settings:
    return Settings(app_env="test", allow_test_providers=True, **values)


def test_txt_anchors_use_real_one_based_line_ranges() -> None:
    content = "\n".join(f"line {number}" for number in range(1, 76)).encode()
    upload = ValidatedUpload("sample.txt", "sample", ".txt", "text/plain", content, "0" * 64)
    anchors = parse_document(upload, _settings())
    assert [anchor.stable_key for anchor in anchors] == ["lines:1-50", "lines:51-75"]
    assert anchors[0].label == "Lines 1-50"


def test_txt_anchors_preserve_leading_blank_and_crlf_source_indices() -> None:
    content = b"\r\n\r\n[LG-POL-999:L001] First retained line\r\nsecond retained line\r\n"
    upload = ValidatedUpload("sample.txt", "sample", ".txt", "text/plain", content, "0" * 64)
    anchors = parse_document(upload, _settings())
    assert len(anchors) == 1
    assert anchors[0].stable_key == "lines:3-4"
    assert anchors[0].label == "Lines 3-4"
    assert anchors[0].text == "[LG-POL-999:L001] First retained line\nsecond retained line"


def test_txt_rejects_null_bytes_and_invalid_utf8() -> None:
    with pytest.raises(UnsafeUploadError, match="binary"):
        _decode_text(b"hello\x00world")
    with pytest.raises(UnsafeUploadError, match="UTF-8"):
        _decode_text(b"\xff\xfe\xfa")


def test_docx_anchors_preserve_heading_and_paragraph_positions() -> None:
    document = Document()
    document.add_heading("Policy", level=1)
    document.add_paragraph("Employees must retain records for seven years.")
    buffer = io.BytesIO()
    document.save(buffer)
    upload = ValidatedUpload(
        "policy.docx",
        "policy",
        ".docx",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        buffer.getvalue(),
        "0" * 64,
    )
    anchors = parse_document(upload, _settings())
    assert anchors[0].stable_key == "section:1:paragraph:1"
    assert anchors[1].stable_key == "section:1:paragraph:2"
    assert anchors[1].label == "Policy: paragraph 2"


def test_docx_zip_bomb_ratio_is_rejected() -> None:
    buffer = io.BytesIO()
    with zipfile.ZipFile(buffer, "w", compression=zipfile.ZIP_DEFLATED) as archive:
        archive.writestr("[Content_Types].xml", "x")
        archive.writestr("word/document.xml", "A" * 100_000)
    with pytest.raises(UnsafeUploadError, match="compression ratio"):
        _validate_docx_archive(buffer.getvalue(), _settings(max_docx_compression_ratio=2.0))


def test_chunk_ids_and_offsets_are_deterministic_and_exact() -> None:
    content = ("Evidence sentence. " * 100).strip()
    upload = ValidatedUpload("e.txt", "e", ".txt", "text/plain", content.encode(), "0" * 64)
    anchors = parse_document(upload, _settings())
    revision_id = uuid.uuid4()
    first = build_chunks(revision_id, anchors, max_characters=300, overlap=40)
    second = build_chunks(revision_id, anchors, max_characters=300, overlap=40)
    assert first == second
    assert len(first) > 1
    for chunk in first:
        anchor = next(item for item in anchors if item.stable_key == chunk.anchor_key)
        assert anchor.text[chunk.start_offset : chunk.end_offset] == chunk.content
